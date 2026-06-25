from collections.abc import Callable
from importlib import resources
import io
import logging

from docx.enum.dml import MSO_COLOR_TYPE
from docx.text.run import Run
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx import Document

from django.http import HttpResponse
from django.utils import translation

from rdmo.projects.exports import Export
from rdmo.projects.models.value import Value
from rdmo.projects.managers import ValueQuerySet

import rdmo_docx_export.exports.templates as templates

logger = logging.getLogger("rdmo."+__name__)

class Style:
    def __init__(self, run: Run):
        self.style = run.style
        self.font = run.font.name
        self.color = run.font.color
        self.size = run.font.size
        self.shadow = run.font.shadow
        self.outline = run.font.outline

        try:
            self.highlight = run.font.highlight_color
        except ValueError:
            self.highlight = None
            
    def apply(self, run: Run) -> Run:
        run.style.base_style = self.style.base_style
        run.style.style_id = self.style.style_id

        run.font.name = self.font
        if self.color.type == MSO_COLOR_TYPE.RGB:
            run.font.color.rgb = self.color.rgb
        elif self.color.type == MSO_COLOR_TYPE.THEME:
            run.font.color.theme_color = self.color.theme_color
        
        if self.highlight is not None:
            run.font.highlight_color = self.highlight
        run.font.size = self.size
        run.font.shadow = self.shadow
        run.font.outline = self.outline

        return run



class _Context(object):
    def __init__(self, datasets: ValueQuerySet, funders: ValueQuerySet, partners: ValueQuerySet):
        self.datasets = datasets
        self.funders = funders
        self.partners = partners

_ParagraphFunction = Callable[['_Context', Paragraph],None]
_Replacements = dict[str, str|_ParagraphFunction|None]


def has_value(query_response: Value | ValueQuerySet | None) -> bool:
    if query_response is None:
        return False
    if isinstance(query_response, ValueQuerySet):
        return len(query_response) > 0
    return len(query_response.value.strip()) > 0

def render_as_list(values: ValueQuerySet) -> str:
    return "\n".join(" * "+v.value for v in values)

class HorizonEuropeDocxExport(Export):

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

    def _stub(self, context: _Context, para: Paragraph) -> None:
        para.add_run("Stub Text. This is not implemented yet.")

    class _dataset_database_value_proxy:
        def __init__(self, parent, field):
            self.parent = parent
            self.field = field
        def __call__(self, context: _Context, para: Paragraph) -> None:
            first = True
            for dataset in context.datasets:
                data_sharing = self.parent.get_value(self.field, set_index=dataset.set_index)

                if not has_value(data_sharing):
                    continue

                if not first:
                    para.add_run("\n").add_break()
                first = False

                headline = para.add_run(f"Dataset {dataset.value}: ")
                headline.italic = True
                para.add_run(data_sharing.value + ".")

    def _dataset_database_value(self, field):
        return self._dataset_database_value_proxy(self, field)

    def _a1a(self, context: _Context, para: Paragraph) -> None:
        """
        Will you re-use any existing data and what will you re-use it for?
        """
        first = True
        for dataset in context.datasets:
            if not first:
                para.add_run("\n\n")
            first = False

            origin = self.get_value("project/dataset/origin", set_index=dataset.set_index)

            headline = para.add_run(f"Dataset {dataset.value}")
            headline.italic = True
            headline.add_break()
            para.add_run(f"This dataset is {origin.value.lower().replace('both (', '').replace(')','')}.\n")
            para.add_run(self.get_text("project/dataset/usage_description", set_index=dataset.set_index))

    def _a2(self, context: _Context, para: Paragraph) -> None:
        """
        What types and formats of data will the project generate or re-use?
        """
        first = True
        for dataset in context.datasets:
            description = self.get_value("project/dataset/description", set_index=dataset.set_index)
            format = self.get_value("project/dataset/format", set_index=dataset.set_index)

            if description or format:
                if not first:
                    para.add_run("\n\n")
                first = False

                headline = para.add_run(f"Dataset {dataset.value}:\n")
                headline.italic = True
            if description:
                para.add_run("The data are ")
                para.add_run(description.value).add_break()
            if format:
                para.add_run("They are provided in the following formats: ")
                para.add_run(format.value)

    def _a4(self, context: _Context, para: Paragraph) -> None:
        """
        What is the expected size of the data that you intend to generate or re-use?
        """
        first = True
        for dataset in context.datasets:
            expected_size = self.get_value("project/dataset/size/volume", set_index=dataset.set_index)

            if not expected_size:
                continue

            if not first:
                para.add_run("\n\n")
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            para.add_run(f"The expected size of the data is {expected_size.value.lower()}.")

    def _a5(self, context: _Context, para: Paragraph) -> None:
        """
        What is the origin/provenance of the data, either generated or re-used?
        """
        first = True
        for dataset in context.datasets:
            provenance_content = self.get_values("project/dataset/provenance/content", set_index=dataset.set_index)
            origin = self.get_value("project/dataset/origin", set_index=dataset.set_index)

            if not provenance_content:
                continue

            if not first:
                para.add_run("\n\n")
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            headline.add_break()
            for v in provenance_content:
                para.add_run(v.text).add_break()

            if "reused" in origin.value.lower():
                author = self.get_value( 'project/dataset/creator/name', set_index=dataset.set_index)
                uri = self.get_value( 'project/dataset/uri', set_index=dataset.set_index)
                para.add_run(f"The data were created by {author.value} and can be found at the following address: {uri.value}")

    def _a6(self, context: _Context, para: Paragraph) -> None:
        """
        To whom might your data be useful ('data utility'), outside your project?
        """
        first = True
        for dataset in context.datasets:
            use_cases = self.get_values("project/dataset/reuse_scenario", set_index=dataset.set_index)

            if not use_cases:
                continue

            if not first:
                para.add_run("\n\n")
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            headline.add_break()
            for usecase in use_cases:
                para.add_run(" * " +usecase.text).add_break()

    def _a7(self, context: _Context, para: Paragraph) -> None:
        """
        Will the data be identified by a persistent identifier?
        """
        first = True
        for dataset in context.datasets:
            has_pid = self.get_bool("project/dataset/pids/yesno", set_index=dataset.set_index)
            pids = self.get_values("project/dataset/pids/system", set_index=dataset.set_index)

            if not has_pid:
                continue

            if not first:
                para.add_run("\n\n")
                first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            headline.add_break()
            para.add_run(render_as_list(pids))

    def _a8(self, context: _Context, para: Paragraph) -> None:
        """
        Will rich metadata be provided to allow discovery? What metadata will be created? What disciplinary or general standards will be followed?
	    In case metadata standards do not exist in your discipline, please outline what type of metadata will be created and how.
        """
        first = True
        for dataset in context.datasets:
            automatic = self.get_value("project/dataset/metadata/creation_automatic", set_index=dataset.set_index)
            semiauto = self.get_value("project/dataset/metadata/creation_semi_automatic", set_index=dataset.set_index)
            manual = self.get_value("project/dataset/metadata/creation_manual", set_index=dataset.set_index)

            if not any([has_value(automatic), has_value(semiauto), has_value(manual)]):
                continue

            if not first:
                para.add_run("").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            headline.add_break()
            if has_value(automatic):
                para.add_run(" * Automatically created: " + automatic.value).add_break()
            if has_value(semiauto):
                para.add_run(" * Automatically created, manually corrected: " + semiauto.value).add_break()
            if has_value(manual):
                para.add_run(" * Manually created: " + manual.value).add_break()


    def _a11(self, context: _Context, para: Paragraph) -> None:
        """
        Will the data be deposited in a trusted repository?
        """
        first = True
        for dataset in context.datasets:
            repositories = self.get_values("project/dataset/preservation/repository", set_index=dataset.set_index)
            trusted = self.get_values("project/dataset/preservation/trusted", set_index=dataset.set_index)

            if len(repositories) == 0:
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            para.add_run("The dataset is stored in a ")
            para.add_run(", ".join(r.value for r in repositories) + ".")

            if len(trusted) > 0:
                para.add_run("\nThis repository is trusted because: ")
                para.add_run(", ".join(t.value for t in trusted) +  ".")

    def _a12(self, context: _Context, para: Paragraph) -> None:
        """
        Have you explored appropriate arrangements with the identified repository where your data will be deposited?
        """
        first = True
        for dataset in context.datasets:
            repository_arrangements = self.get_values("project/dataset/preservation/repository_arrangements", set_index=dataset.set_index)

            if len(repository_arrangements) == 0:
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            para.add_run(", ".join(r.value for r in repository_arrangements) + ".")


    def _a13b(self, context: _Context, para: Paragraph) -> None:
        """
        Will the repository resolve the identifier to a digital object?
        """
        first = True
        for dataset in context.datasets:
            resolver = self.get_value("project/dataset/pids/resolver", set_index=dataset.set_index)

            if not has_value(resolver):
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            para.add_run(resolver.value)
            if "no" not in resolver.value.lower():
                para.add_run(", the repository will resolve the identifier to a digital object")
            para.add_run(".")

    def _a14b(self, context: _Context, para: Paragraph) -> None:
        """
        Will the repository resolve the identifier to a digital object?
        """
        first = True
        for dataset in context.datasets:
            explanation = self.get_values("project/dataset/sharing/explanation", set_index=dataset.set_index)

            explanation = [ e for e in explanation if has_value(e) ]

            if len(explanation) == 0:
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}:\n")
            headline.italic = True
            for e in explanation:
                para.add_run(e.value).add_break()

    def _a15(self, context: _Context, para: Paragraph) -> None:
        """
        If an embargo is applied to give time to publish or seek protection of the intellectual property (e.g. patents),
	    specify why and how long this will apply, bearing in mind that research data should be made available as soon as possible.
        """
        first = True
        for dataset in context.datasets:
            embargo_perioods = self.get_values("project/dataset/preservation/embargo_period", set_index=dataset.set_index)

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True

            if len(embargo_perioods) == 0:
                para.add_run("No embargo is applied.")
                continue

            para.add_run("An embargo is applied, for the ")
            para.add_run(", ".join(r.value for r in embargo_perioods) + ".")

    def _a19(self, context: _Context, para: Paragraph) -> None:
        """
        Is there a need for a data access committee (e.g. to evaluate/approve access requests to personal/sensitive data)?
        """
        comittee = self.get_value('project/legal_aspects/official_approval/data_access_committee')

        if has_value(comittee) and "yes" in comittee.value.lower():
            para.add_run(
                "This consortium will have a Data Access Committee. Their remit will be to select the"
                " data that will be openly accessible on a case-by-case basis. Ethical aspects and data security,"
                " including intellectual property requirements, will be considered as will access requests to"
                " personal/sensitive data. If necessary, some or all of a potential publication's data will be"
                " withheld. This will be decided in consultation with the relevant partner(s).")
        elif has_value(comittee) and "no" in comittee.value.lower():
            para.add_run(
                "This consortium will not have a Data Access Committee, because the project is not"
                " going to produce sensitive data. All results will be publicly available without restrictions.")
        else:
            para.add_run(
                "This consortium has not established a Data Access Committee. The appointed data"
                " responsible / corresponding author will decide alone about granting access to the data.")

    def _a24(self, context: _Context, para: Paragraph) -> None:
        """
        In case it is unavoidable that you use uncommon or generate project specific ontologies or vocabularies,
        will you provide mappings to more commonly used ontologies?
	    Will you openly publish the generated ontologies or vocabularies to allow reusing, refining or extending them?
        """
        first = True
        for dataset in context.datasets:
            mappings = self.get_value('project/dataset/metadata/mappings', set_index=dataset.set_index)
            vocabularies = self.get_value('project/dataset/metadata/vocabularies_open', set_index=dataset.set_index)

            if not has_value(mappings):
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            para.add_run(mappings.value)

            if has_value(vocabularies):
                para.add_run("\n"+vocabularies.value)

    def _a26(self, context: _Context, para: Paragraph) -> None:
        """
        How will you provide documentation needed to validate data analysis and facilitate data re-use
    	(e.g. readme files with information on methodology, codebooks, data cleaning, analyses, variable
        definitions, units of measurement, etc.)?
        """
        first = True
        for dataset in context.datasets:
            documentation_where = self.get_value('project/dataset/documentation/where', set_index=dataset.set_index)
            documentations = self.get_values('project/dataset/documentation', set_index=dataset.set_index)

            if not has_value(documentation_where):
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            para.add_run(f"We will provide documentation { documentation_where.value.lower() }, in the form of:\n")
            para.add_run(render_as_list(documentations))


    def _a27b(self, context: _Context, para: Paragraph) -> None:
        """
        Will your data be licensed using standard reuse licenses, in line with the obligations set out in the Grant Agreement?
        """
        first = True
        for dataset in context.datasets:
            sharing_conditions = self.get_value('project/dataset/sharing/conditions', set_index=dataset.set_index)

            if not has_value(sharing_conditions):
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            para.add_run(f"Yes, with the following license: { sharing_conditions.value }.")

    def _a30(self, context: _Context, para: Paragraph) -> None:
        """
        Describe all relevant data quality assurance processes.
        """
        first = True
        for dataset in context.datasets:
            qa = self.get_values('project/dataset/quality_assurance', set_index=dataset.set_index)

            if not has_value(qa):
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            para.add_run(render_as_list(qa))

    def _a32(self, context: _Context, para: Paragraph) -> None:
        """
        In addition to the management of data, beneficiaries should also consider and plan for
        the management of other research outputs that may be generated or re-used throughout their
        projects.
    	Such outputs can be either digital (e.g. software, workflows, protocols, models, etc.) or
        physical (<i>e.g.</i> new materials, antibodies, reagents, samples, etc.).
        """
        outputs = self.get_values('project/other_research_output/plan')

        para.add_run("The project will produce research output also of the following types:\n")
        para.add_run(render_as_list(outputs))

    def _a33(self, context: _Context, para: Paragraph) -> None:
        """
        Beneficiaries should consider which of the questions pertaining to FAIR data above can apply
        to the management of other research outputs, and should strive to provide sufficient detail
        on how their research outputs will be managed and shared, or made available for re-use, in
        line with the FAIR principles.
        """
        outputs = self.get_values('project/other_research_output/fair')

        para.add_run("The following aspects are relevant and will be considered within the project:\n")
        para.add_run(render_as_list(outputs))


    def _a34(self, context: _Context, para: Paragraph) -> None:
        """
        What will the costs be for making data or other research outputs FAIR in your project
	    (e.g. direct and indirect costs related to storage, archiving, re-use, security, etc.)?
        """
        def render_value(identifier:str)->str:
            v = self.get_value(identifier)
            return v if v is not None else ""

        para.add_run(f"For metadata: { render_value('project/costs/metadata/personnel') }; { render_value ('project/costs/metadata/non_personnel') }\n")
        para.add_run(f"For PIDs: { render_value ('project/costs/pid/personnel') }; { render_value ('project/costs/pid/non_personnel') }\n")
        para.add_run(f"For data preservation: { render_value ('project/costs/preservation/personnel') }; { render_value ('project/costs/preservation/non_personnel')} \n")
        para.add_run(f"For management of other research outputs: { render_value ('project/costs/other_research_output/personnel') }; { render_value ('project/costs/other_research_output/non_personnel') }")

    def _a36(self, context: _Context, para: Paragraph) -> None:
        """
        Describe all relevant data quality assurance processes.
        """
        first = True
        for partner in context.partners:
            responsible = self.get_values('project/partner/contact_person/name', set_index=partner.set_index)

            if not has_value(partner):
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"{partner.value}: ")
            headline.italic = True
            para.add_run(", ".join(r.value for r in responsible))

    def _a37(self, context: _Context, para: Paragraph) -> None:
        """
        How will long term preservation be ensured? Discuss the necessary resources to accomplish this
	    (costs and potential value, who decides and how, what data will be kept and for how long).
        """
        part_duration = para.add_run("Duration of preservation\n")
        part_duration.bold = True
        for dataset in context.datasets:
            duration = self.get_value('project/dataset/preservation/duration', set_index=dataset.set_index)

            if not has_value(duration):
                continue

            headline = para.add_run(f"{dataset.value}: ")
            headline.italic = True
            para.add_run(duration.value + "\n")
        para.add_run("\n")

        part_costs = para.add_run("Costs\n")
        part_costs.bold = True
        para.add_run("See question 34.\n\n")

        part_Potential_value = para.add_run("Potential value\n")
        part_Potential_value.bold = True
        para.add_run("See question 6.\n\n")

        responsible = self.get_value('project/preservation/responsible_person/name')
        if has_value(responsible):
            part_responsible = para.add_run("Preservation decisions\n")
            part_responsible.bold = True
            para.add_run(f"All decision related to preservation fall under the responsibility of { responsible.value }.\n\n")

        criteria = self.get_values('project/preservation/selection_criteria')
        part_criteria = para.add_run("Criteria for preservation\n")
        part_criteria.bold = True
        para.add_run(", ".join(c.value for c in criteria))
        para.add_run("\n\n")

        part_motivation = para.add_run("Motivation for long-term preservation\n")
        part_motivation.bold = True

        first = True
        for dataset in context.datasets:
            purpose = self.get_value('project/dataset/preservation/purpose', set_index=dataset.set_index)

            if not has_value(purpose):
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"{dataset.value}: ")
            headline.italic = True
            para.add_run(purpose.value)

    def _a38(self, context: _Context, para: Paragraph) -> None:
        """
        What provisions are or will be in place for data security (including data recovery as well as secure storage/archiving and transfer of sensitive data)?
        """
        first = True
        for dataset in context.datasets:
            security = self.get_values('project/dataset/data_security/security_measures', set_index=dataset.set_index)

            if not has_value(security):
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            para.add_run(render_as_list(security))

    def _a39(self, context: _Context, para: Paragraph) -> None:
        """
        Will the data be safely stored in trusted repositories for long-term preservation and curation?
        """
        para.add_run("Yes, the chosen repositories are trustworthy. See question 11 for further details.\n\n")

        first = True
        for dataset in context.datasets:
            security = self.get_values('project/dataset/preservation/repository', set_index=dataset.set_index)
            certified = self.get_bool('project/dataset/preservation/certification', set_index=dataset.set_index)

            if not has_value(security):
                continue

            if not first:
                para.add_run("\n").add_break()
            first = False

            headline = para.add_run(f"Dataset {dataset.value}: ")
            headline.italic = True
            para.add_run(", ".join(s.value for s in security))
            if certified is not None:
                para.add_run("(certified)" if certified.value else "(not certified)")

    def _a40(self, context: _Context, para: Paragraph) -> None:
        """
        Are there, or could there be, any ethics or legal issues that can have an impact on data sharing?
        """
        part_personal = para.add_run("Personal data")
        part_personal.bold = True

        for dataset in context.datasets:
            sensitive = self.get_bool('project/dataset/sensitive_data/personal_data_yesno/yesno', set_index=dataset.set_index)

            if sensitive is None or not sensitive:
                continue

            identifying = self.get_bool('project/dataset/sensitive_data/personal_data/bdsg_3_9', set_index=dataset.set_index)
            anonymization = self.get_value('project/dataset/sensitive_data/personal_data/anonymization', set_index=dataset.set_index)
            deletion = self.get_value('project/dataset/sensitive_data/personal_data/deletion', set_index=dataset.set_index)

            headline = para.add_run(f"\nDataset {dataset.value}: ")
            headline.italic = True
            para.add_run("The dataset contains personal data.")
            if identifying is not None and identifying:
                para.add_run("\nSome data contain information allowing person identification.")
            if has_value(anonymization):
                para.add_run(f"\nThe personal data will be anonymised/pseudonymised: {anonymization.value}.")
            if has_value(deletion):
                para.add_run(f"\nThe personal data will be safely deleted {deletion.value}.")

        part_other_sensitive = para.add_run("\n\nOther sensitive data")
        part_other_sensitive.bold = True

        for dataset in context.datasets:
            sensitive = self.get_value('project/dataset/sensitive_data/other/description', set_index=dataset.set_index)

            if not has_value(sensitive):
                continue

            headline = para.add_run(f"\nDataset {dataset.value}: ")
            headline.italic = True
            para.add_run(sensitive.value)            

        part_approval = para.add_run("\n\nOfficial approval")
        part_approval.bold = True
        committee = self.get_value('project/legal_aspects/official_approval/ethics_committee')
        if has_value(committee):
            para.add_run(f"\nThe research has been { committee.value }.")

        statutatory= self.get_bool('project/legal_aspects/official_approval/statutatory_approval/yesno')
        if has_value(statutatory) and statutatory.value:
            title = self.get_value('project/legal_aspects/official_approval/statutatory_approval/title')
            agency = self.get_value('project/legal_aspects/official_approval/statutatory_approval/agency')
            if has_value(title) or has_value(agency):
                para.add_run("\nThe research has been approved")
                if has_value(title):
                    para.add_run(f" under the title {title.value}")
                if has_value(agency):
                    para.add_run(f" by the following agency: {agency.value}")
                para.add_run(".")

        part_ipr = para.add_run("\n\nIntellectual property rights")
        part_ipr.bold = True

        first = True
        for dataset in context.datasets:
            sensitive = self.get_bool('project/legal_aspects/ipr/yesno', set_index=dataset.set_index)

            if not has_value(sensitive) or not sensitive:
                continue

            if not first:
                para.add_run("\n")
            first = False

            copyright = self.get_bool('project/legal_aspects/ipr/copyrights', set_index=dataset.set_index)
            other_right = self.get_bool('project/legal_aspects/ipr/other_rights', set_index=dataset.set_index)

            headline = para.add_run(f"\nDataset {dataset.value}:")
            headline.italic = True
            if has_value(copyright):
                para.add_run(f"\n{copyright.value}")
            if has_value(other_right):
                para.add_run(f"\n{other_right.value}")

        para.add_run("\n\nSee question 17 and the following question for further details.")


    def _replace_paragraph_contents(self, replacements: _Replacements, context: _Context, para: Paragraph):
        """
        Checks if a paragraph's content is to be replaced.

        If yes: Replaces contents. Tries to keep style intact.

        Functional Style elements (i.e. italic, bold) may be applied by content
        functions and will not be overwritten afterwards.
        Then some style elements are not yet copied correctly (Shadow, Outline, Theme).
        They are not readable/writable with python-docx, but we need to
        create 'runs' to apply functional style.
        """
        def replace_para(para, text):
            for run in para.runs:
                run.text = ""
            para.runs[0].text = text


        if para.text.startswith("{{") and para.text.endswith("}}"):
            if para.text in replacements:
                value = replacements[para.text]
                if value is None:
                    replace_para(para, "***error: value not defined***")
                elif isinstance(value, str):
                    replace_para(para,value)
                else:
                    style = Style(para.runs[0])
                    para.text = ""
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    value(context, para)
                    for run in para.runs:
                        style.apply(run)
            else:
                replace_para(para, "Lorem Ipsum... ***replacement sequence not defined***")

    def render(self):
        logger.info("Generating Docx Document...")
        logger.info("TEST logging Umlaute...")
        logger.info("ÄÖÜ...".encode('unicode_escape').decode('ascii'))
        logger.info(f"{'äöü'}")
        logger.info("TEST done.")

        template = resources.files(templates) / "horizon-template.docx"
        doc = Document(template.open("rb"))
        with translation.override("en"):
            context = _Context(
                datasets = self.get_set("project/dataset/id"),
                partners = self.get_set("project/partner/id"),
                funders = self.get_set("project/funder/id"),
            )

            replacements = {
                "{{projectnumber}}" : self.get_text("project/funder/grant_nr"),
                "{{projectacronym}}": self.get_text("project/acronym"),
                "{{projecttitle}}"  : self.get_text("project/title"),
                "{{dmpdate}}"       : self.get_text("project/dmp/dmp_date"),
                "{{dmpversion}}"    : self.get_text("project/dmp/dmp_version"),
                "{{Answer01a}}"     : self._a1a,
                "{{Answer01b}}"     : self._dataset_database_value("project/dataset/reuse_existing"),
                "{{Answer02}}"      : self._a2,
                "{{Answer03}}"      : self._dataset_database_value("project/dataset/usage_description"),
                "{{Answer04}}"      : self._a4,
                "{{Answer05}}"      : self._a5,
                "{{Answer06}}"      : self._a6,
                "{{Answer07}}"      : self._a7,
                "{{Answer08}}"      : self._a8,
                "{{Answer09}}"      : self._dataset_database_value("project/dataset/metadata/search_keywords"),
                "{{Answer10}}"      : self._dataset_database_value("project/dataset/metadata/harvesting"),
                "{{Answer11}}"      : self._a11,
                "{{Answer12}}"      : self._a12,
                "{{Answer13a}}"     : self._dataset_database_value("project/dataset/pids/yesno"),
                "{{Answer13b}}"     : self._a13b,
                "{{Answer14a}}"     : self._dataset_database_value('project/dataset/sharing/yesno'),
                "{{Answer14b}}"     : self._a14b,
                "{{Answer15}}"      : self._a15,
                "{{Answer16}}"      : self._dataset_database_value('project/dataset/sharing/conditions'),
                "{{Answer17}}"      : self._dataset_database_value('project/dataset/sharing/restrictions_explanation'),
                "{{Answer18}}"      : self._dataset_database_value('project/dataset/preservation/access_authentication'),
                "{{Answer19}}"      : self._a19,
                "{{Answer20a}}"     : self._dataset_database_value('project/dataset/metadata/license_for_metadata'),
                "{{Answer20b}}"     : self._dataset_database_value('project/dataset/metadata/access_info'),
                "{{Answer21a}}"     : self._dataset_database_value("project/dataset/preservation/reuse_duration"),
                "{{Answer21b}}"     : self._dataset_database_value("project/dataset/metadata/available_without_data"),
                "{{Answer22}}"      : self._dataset_database_value('project/dataset/software_documentation'),
                "{{Answer23a}}"     : self._dataset_database_value("project/dataset/metadata/standards"),
                "{{Answer23b}}"     : self._dataset_database_value("project/dataset/interoperability"),
                "{{Answer24}}"      : self._a24,
                "{{Answer25}}"      : self._dataset_database_value("project/dataset/metadata/references_to_other_data"),
                "{{Answer26}}"      : self._a26,
                "{{Answer27a}}"     : self._dataset_database_value("project/dataset/sharing/yesno"),
                "{{Answer27b}}"     : self._a27b,
                "{{Answer28}}"      : "The data will be available for re-use, as far as reported in question 21.",
                "{{Answer29}}"      : self._dataset_database_value("project/dataset/provenance/standards"),
                "{{Answer30}}"      : self._a30,
                "{{Answer31}}"      : "We will address these aspects in the following section",
                "{{Answer32}}"      : self._a32,
                "{{Answer33}}"      : self._a33,
                "{{Answer34}}"      : self._a34,
                "{{Answer35}}"      : self.get_text('project/costs/preservation/cover_how'),
                "{{Answer36}}"      : self._a36,
                "{{Answer37}}"      : self._a37,
                "{{Answer38}}"      : self._a38,
                "{{Answer39}}"      : self._a39,
                "{{Answer40}}"      : self._a40,
                "{{Answer41}}"      : self._stub,
                "{{Answer42}}"      : self._stub,
            }

            for para in doc.paragraphs:
                self._replace_paragraph_contents(replacements, context, para)
            for tab in doc.tables:
                for c in tab.columns:
                    for cell in c.cells:
                        for para in cell.paragraphs:
                            self._replace_paragraph_contents(replacements, context, para)

        response_data = io.BytesIO()
        doc.save(response_data)
        response_data.seek(0)  # Rewind file pointer to extract generated data!

        response = HttpResponse(response_data.read(), content_type='application/docx')
        response['Content-Disposition'] = 'attachment; filename="horizon-europe-export.docx"'
        return response

    def submit(self):
        raise NotImplementedError
